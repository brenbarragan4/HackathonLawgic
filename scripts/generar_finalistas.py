# -*- coding: utf-8 -*-
"""Genera un Word con el resumen de los proyectos finalistas del hackathon."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VERDE  = RGBColor(0x14, 0x3A, 0x2E)
TINTA  = RGBColor(0x2B, 0x33, 0x2D)
GRIS   = RGBColor(0x5A, 0x60, 0x57)
DORADO = RGBColor(0xA8, 0x82, 0x45)

doc = Document()

# Margenes
for s in doc.sections:
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

# Fuente base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = TINTA

def parrafo(runs, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align: p.alignment = align
    for (t, size, color, bold, italic, font) in runs:
        r = p.add_run(t)
        r.font.size = Pt(size); r.font.color.rgb = color
        r.bold = bold; r.italic = italic; r.font.name = font
    return p

def R(t, size=11, color=TINTA, bold=False, italic=False, font='Calibri'):
    return (t, size, color, bold, italic, font)

# ---- Portada ----
parrafo([R("MINI-HACKATHON LAWGIC · CLAUDE PARA ABOGADOS · 2026", 10, DORADO, True)], space_after=2)
parrafo([R("Proyectos Finalistas", 28, VERDE, True, font='Georgia')], space_after=2)
parrafo([R("Resumen de los 8 proyectos seleccionados en la ronda de pitch — para estudiarlos e inspirar la siguiente versión.", 12, GRIS, italic=True)], space_after=14)

# ---- Lista de finalistas ----
finalistas = [
    ("1", "Esteban Echeverría", "Calculadora de Plazos",
     "Calcula vencimientos procesales (recurso de revocación, revisión, juicio de nulidad, amparo) aplicando la regla de surtimiento de cada ley más los días inhábiles del SAT/TFJA. Publicado en Netlify."),
    ("2", "Hazur Socconini", "Protege Tu Invento",
     "Estudio de viabilidad de marca: motor legal LFPPI + consulta a MARCANET en vivo (ingeniería inversa del IMPI) + disponibilidad de dominios y redes, todo en un reporte PDF."),
    ("3", "Kathia Holguín", "Lindero",
     "Lee escrituras (incluso escaneadas, con OCR) y genera con IA un dictamen de due diligence inmobiliaria: inmueble, comparecientes, poderes y riesgos, exportable a Word."),
    ("4", "Víctor Manuel Varela", "lex-mcp",
     "Servidor MCP que convierte Claude Desktop en un “abogado senior”: 6 herramientas (generar, revisar y resumir contratos, calcular plazos…) con contratos fundamentados en CCF/LFPPI/LFDA."),
    ("5", "Yasmin Yektajo", "CaboClose",
     "Portal que centraliza el cierre de fideicomiso en zona restringida (Los Cabos): KYC/PLD con beneficiario controlador conforme a la reforma LFPIORPI 2025, checklists dinámicos y portal del cliente."),
    ("6", "Luis Enrique Maass", "Onboarding de expediente PJF",
     "Claude opera Chrome en vivo: descarga las actuaciones del Amparo Directo 805/2025 desde el portal judicial, las consolida en un PDF de 29 páginas con carátula e índice más un estudio del caso. Validado contra un expediente federal real."),
    ("7", "Diego Humberto Cruz", "Armetis Legal Valuador",
     "Valúa el catálogo de un artista conectándose vía API de Spotify para revisar contratos musicales según las reproducciones reales (cruce de propiedad intelectual + finanzas + música)."),
    ("8", "Jorge Luis Herrera Alor", "JusticIA Familiar",
     "Genera un borrador de sentencia de oralidad familiar (Quintana Roo) fundado en ley local + jurisprudencia real de la SCJN con número de registro + tratados; con control de congruencia y datos anonimizados."),
]

for num, autor, proyecto, desc in finalistas:
    parrafo([R(num + ".  ", 14, DORADO, True, font='Georgia'),
             R(autor, 14, VERDE, True, font='Georgia'),
             R("   —   ", 12, GRIS),
             R(proyecto, 12, DORADO, True)], space_after=2, space_before=8)
    parrafo([R(desc, 11, TINTA)], space_after=4)

# ---- El comun denominador ----
parrafo([R("El común denominador: ¿por qué quedaron finalistas?", 16, VERDE, True, font='Georgia')], space_before=18, space_after=2)
parrafo([R("Analizando las pantallas de los 8 proyectos, se repiten estos patrones. Ese fue, muy probablemente, el criterio que el jurado premió:", 11, GRIS, italic=True)], space_after=8)

patrones = [
    ("Nicho híper-específico y profundo",
     "Cada proyecto resolvía UN problema muy concreto del día a día legal (plazos procesales, viabilidad de marca, due diligence de escrituras, cierre de fideicomisos, expedientes del PJF, valuación de catálogos, sentencias familiares). No eran herramientas generales, sino especializadas."),
    ("Fundamento jurídico real y citado",
     "Citaban leyes y artículos específicos (CCF, LFPPI, LFDA, LFPIORPI 2025, jurisprudencia de la SCJN con número de registro). Ese rigor legal genera mucha credibilidad ante un jurado de abogados."),
    ("Conexión con datos y fuentes reales en vivo",
     "Se conectaban a fuentes verdaderas: MARCANET (IMPI), API de Spotify, el portal del Poder Judicial, OCR de escrituras escaneadas. No eran demos vacías: jalaban información real."),
    ("Uso avanzado de Claude (MCP, Claude operando el navegador)",
     "Varios usaron Claude más allá del chat: servidores MCP que vuelven a Claude Desktop un 'abogado senior', Claude operando Chrome en vivo, o conectado a notas en Obsidian."),
    ("Entregable profesional y tangible",
     "Todos terminaban en algo descargable y útil para el despacho: un dictamen, un reporte PDF, una sentencia o un expediente consolidado, listos para usar."),
    ("Flujo completo con poco input del usuario",
     "Automatizaban el trabajo pesado (descargar, consolidar, analizar, redactar) para que el abogado teclee lo mínimo. La herramienta hace casi todo."),
    ("Validado contra casos reales",
     "Probaban su solución con expedientes y casos reales (por ejemplo, un amparo directo real), no con ejemplos inventados."),
]
for titulo, d in patrones:
    parrafo([R("•  ", 11, DORADO, True),
             R(titulo + ":  ", 11, VERDE, True),
             R(d, 11, TINTA)], space_after=4)

parrafo([R("En una frase:  ", 11, VERDE, True),
         R("ganó lo súper especializado, fundamentado en la ley real, conectado a datos reales y con un entregable profesional — usando Claude de forma avanzada.", 11, TINTA, italic=True)], space_before=6, space_after=4)

# ---- Cierre: ideas para mi app ----
parrafo([R("Ideas para la siguiente versión de MI app", 16, VERDE, True, font='Georgia')], space_before=16, space_after=4)
ideas = [
    ("Especializar el nicho", "Casi todos los finalistas resolvían UN problema muy concreto (plazos, marcas, escrituras, fideicomisos). Elegir una sola especialidad fuerte puede pegar más que algo general."),
    ("Conectar con Obsidian + Claude (MCP)", "Varios conectaron Claude con sus notas/plantillas de contratos. Vale la pena explorar el MCP para que mi app lea y escriba en una base de conocimiento legal."),
    ("Menos input manual, más automático", "Que la herramienta haga más sola (leer archivos, descargar del portal, precargar datos) y el usuario teclee lo mínimo."),
    ("Mi diferenciador: la voz", "Mi app ya permite dictar por voz con un asistente (Lexi). Es una ventaja real; conviene potenciarla y mostrarla mejor en el pitch."),
    ("Fundamentar en la ley", "Citar artículos y jurisprudencia con número de registro da mucha credibilidad ante un jurado de abogados."),
]
for titulo, d in ideas:
    parrafo([R("•  ", 11, DORADO, True),
             R(titulo + ":  ", 11, VERDE, True),
             R(d, 11, TINTA)], space_after=4)

parrafo([R("Tu proyecto (Revisa tu Contrato) sigue en línea y es tuyo: ", 10, GRIS, italic=True),
         R("hackathonlawgic.vercel.app", 10, DORADO, True, italic=True)], space_before=14)

out = r"C:\Users\brenb\OneDrive\Desktop\Proyectos-Finalistas-Lawgic.docx"
doc.save(out)
print("GUARDADO:", out)
